
import crewai
import crewai_tools
from groq import Groq
from crewai import Agent
from langchain.llms import OpenAI
from crewai import Agent, Task, Crew, Process

import string
from pypdf import PdfReader

from langchain_groq import ChatGroq


from docx import Document

from langchain_community.tools import DuckDuckGoSearchRun
from langchain.agents import load_tools
from langchain.tools import tool
from crewai.tasks.task_output import TaskOutput

## api keys
Groq_api_key = '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'

api_key = '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'
os.environ["GROQ_API_KEY"] = api_key

## the LLM

llm = ChatGroq(model="llama3-8b-8192",api_key=Groq_api_key)

## web scrapper
import requests
import bs4


def openup(name):
    topic = name
    topic = topic.replace(' ','+')
    print(topic)
    return topic

def parse(name):
    topic = openup(name)
    link = 'https://scholar.google.com/scholar?hl=en&as_sdt=0%2C5&q='+topic+'&btnG='
    res = requests.get(link)
    res.raise_for_status()
    soup = bs4.BeautifulSoup(res.text)

    a = 'gs_r gs_or gs_scl'
    b= 'gs_res_ccl_mid'
    all_ = soup.select('#'+b)
    some = all_[0].select('.gs_or')


    all_links = {}

    for i in range(10):
        topic = some[i].select('div > h3')[0].select('a')[0].getText()

        print(topic)
        try:

            mlink = some[i].select('.gs_or_ggsm')[0].select('a')[0].get('href')
    #         nam = some[i].select('div')[0].select('div')[0].select('div')[0].text
    #         print(nam)
            print(mlink)

        except:
            print('no topic for ', topic)
            mlink = ''

        all_links[topic] = mlink

    return all_links
# returns a lst that contains the names and links

def Sichub2(title):
    payload = { 'request': title}
    link = 'https://sci-hub.st/'
    # link = 'https://sci-hub.st/10.1109/icmla.2017.0-179'

    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_10_1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36'}


    a = requests.post(link, data=payload, headers = headers)

    soup = bs4.BeautifulSoup(a.text)
    content = soup.select('embed')[0].get('src').replace('#navpanes=0&view=FitH','').replace('//','/')
    contents= 'https:/' + content
    r = requests.get(contents, stream=True)

    title = title.replace(':','')
    with open(title +'.pdf', 'wb') as file:
        file.write(r.content)

    print('DOne downloading: ', title )
    return title+'.pdf'


def google(title, url):
    dashboardFile = requests.get(url, allow_redirects=True)

    for i in string.punctuation:
        title = title.replace(i,'')

    title = title +'.pdf'
    open(title, 'wb').write(dashboardFile.content)
    print('google:', title)
    reader = PdfReader(title)
    return title


#AVX

down = []
def Avx():
  topic = openup()
  cat = classify(topic)
  link = 'https://arxiv.org/search/' + cat+ '?query=' + topic + '&searchtype=all&abstracts=show&order=-announced_date_first&size=50'
  res = requests.get(link)
  res.raise_for_status()
  soup = bs4.BeautifulSoup(res.text)
  a = soup.select('.arxiv-result')

  for i in range(10):
    b = a[i].select('p')
    topic = b[1].text
    url = a[i].select('span')[0].select('a')[0].get('href')
    topic = topic.replace('\n','')
    topic = topic.replace(' ','')
    topic = topic.replace("/","")


    dashboardFile = requests.get(url, allow_redirects=True)
    title = topic.replace(':','')
    name = topic+'.pdf'
    open(name, 'wb').write(dashboardFile.content)
    down.append(name)
    print('Done downloading:', title)

def classify(topic):
  from groq import Groq

  api_key = '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'

  client = Groq(api_key = api_key)
  completion = client.chat.completions.create(
      model="llama3-8b-8192",
      messages=[
          {
              "role": "user",
              "content": """ heyyy calssify this topic " """ +topic+ """ " appropriately and respond if it is related to the subjects: physics or computer Science or Mathematics or
              Quantitative Biology or Statistics or Electrical Engineering and Systems Science or Economics or Quantitative Finance
              limt your answer to just one subject
              give a very short output
              eg:
              output: physics"""
          }
      ],
      temperature=1,
      max_tokens=1024,
      top_p=1,
      stream=True,
      stop=None,
  )

  you = []
  for chunk in completion:
      you.append((chunk.choices[0].delta.content or ""))
  m = ''
  for j in you:
    m = m + j
  print(m)

  subjects = {'Physics':'physics',
 'Computer Science':'cs',
 'Mathematics':'math',
 'Quantitative Biology':'q-bio',
 'Statistics':'stat',
 'Electrical Engineering and Systems Science':'eess',
 'Economics':'econ',
 'Quantitative Finance':'q-fin',

 }
  for i in subjects.keys():
    if i in m:
      a = subjects[i]
      print(a)
  return (a)


# the tools 

from crewai_tools import PDFSearchTool
def ragtool(path):

  rag_tool = PDFSearchTool(pdf=path,
      config=dict(
          llm=dict(
              provider="groq", # or google, openai, anthropic, llama2, ...
              config=dict(
                  model="llama3-8b-8192",
                  # temperature=0.5,
                  # top_p=1,
                  # stream=true,
              ),
          ),
          embedder=dict(
              provider="huggingface", # or openai, ollama, ...
              config=dict(
                  model="BAAI/bge-small-en-v1.5",
                  #task_type="retrieval_document",
                  # title="Embeddings",
              ),
          ),
      )
  )
  return rag_tool


##  The agent workflow
def agents():


    search_tool = DuckDuckGoSearchRun()

    # Define the topic of interest
    topic = 'AI in healthcare'

    # Loading Human Tools
    human_tools = load_tools(["human"])

    def callback_function(output: TaskOutput):
        # Do something after the task is completed
        # Example: Send an email to the manager
        print(f"""
            Task completed!
            Task: {output.description}
            Output: {output.result}
        """)

    # Creating custom tools
    class ContentTools:
        @tool("Read webpage content")
        def read_content(url: str) -> str:
            """Read content from a webpage."""
            response = requests.get(url)
            soup = BeautifulSoup(response.content, 'html.parser')
            text_content = soup.get_text()
            return text_content[:5000]

    # Define the manager agent
    manager_agent = Agent(
        role='Project Manager',
        goal='Coordinate the project to ensure a seamless integration of research findings into compelling narratives',
        verbose=True,
        backstory="""With a strategic mindset and a knack for leadership, you excel at guiding teams towards
        their goals, ensuring projects not only meet but exceed expectations.""",
        allow_delegation=True,
        max_iter=10,
        max_rpm=20,
        llm=llm
    )

    # Define the senior researcher agent
    researcher = Agent(
        role='Senior Researcher',
        goal=f'Uncover groundbreaking technologies around {topic}',
        verbose=True,
        backstory="""Driven by curiosity, you're at the forefront of innovation, eager to explore and share
        knowledge that could change the world.""",
        llm=llm
    )

    # Define the writer agent
    writer = Agent(
        role='Writer',
        goal=f'Narrate compelling tech stories around {topic}',
        verbose=True,
        backstory="""With a flair for simplifying complex topics, you craft engaging narratives that captivate
        and educate, bringing new discoveries to light in an accessible manner.""",
        llm=llm
    )

    # Define the asynchronous research tasks
    list_ideas = Task(
        description="List of 5 interesting ideas to explore for an article about {topic}.",
        expected_output="Bullet point list of 5 ideas for an article.",
        tools=[search_tool, ContentTools().read_content],
        agent=researcher,
        async_execution=True
    )

    list_important_history = Task(
        description="Research the history of {topic} and identify the 5 most important events.",
        expected_output="Bullet point list of 5 important events.",
        tools=[search_tool, ContentTools().read_content],
        agent=researcher,
        async_execution=True
    )

    # Define the writing task that waits for the outputs of the two research tasks
    write_article = Task(
        description=f"Compose an insightful article on {topic}, including its history and the latest interesting ideas.",
        expected_output="A 4 paragraph article about AI in healthcare.",
        tools=[search_tool, ContentTools().read_content],
        agent=writer,
        context=[list_ideas, list_important_history],  # Depends on the completion of the two asynchronous tasks
        callback=callback_function
    )

    # Define the manager's coordination task
    manager_task = Task(
        description=f"""Oversee the integration of research findings and narrative development to produce a final comprehensive
        report on {topic}. Ensure the research is accurately represented and the narrative is engaging and informative.""",
        expected_output=f'A final comprehensive report that combines the research findings and narrative on {topic}.',
        agent=manager
    )

    # Forming the crew with a hierarchical process including the manager
    crew = Crew(
        agents=[researcher, writer],
        tasks=[list_ideas, list_important_history, write_article, manager_task],
        process=Process.hierarchical,
        manager_agent=manager_agent,
        manager_llm=llm,
    )

    # Kick off the crew's work
    results = crew.kickoff()

    # Print the results
    return results
you = []   

def run(name):
    data = parse(name)

    err = []


    for i in data.keys():
        try:
            a = Sichub2(i)
            down.append(a)
        except:
            err.append(i)

    for i in err:
        try:
            b = google(i,data[i])
            down.append(b)
        except:
            print(i)

       
    for i in down:
        reader = PdfReader(i)
        print('my head')
        try:
          global rag_rool
          rag_tool = ragtool(i)
          print('my nech')
          words = agents()
          print('my hand')
          you.append(words)
        except Exception as exc:
            print(exc)
          
    return rag_tool
            

    doc()
    download_pdf(pdf_path)

        
from docx import Document

def doc():
    documents = Document()
    file_name = 'Summary'
    documents.add_heading('AI Summary', 0)
    for i in range(len(you)):
      documents.add_heading(down[i].replace('/content/',''), level=1)
      documents.add_paragraph(you[i].raw)

    documents.save(file_name + '.docx')

import streamlit as st  # pip install streamlit


st.header(":mailbox: Hello researcher!")
topic = st.text_input(label="Research Topic", placeholder="what are you working on?")
into = st.text_input(label="Additional info.", placeholder="what is it about?")





import streamlit as st
import os

def download_pdf(pdf_path):
    """your files woule be ready soon, wait for another go button ready"""
    
    
    """your files are ready"""
    if not os.path.exists(pdf_path):
        st.error(f"PDF file '{pdf_path}' not found.")
        return

    with open(pdf_path, 'rb') as f:
        pdf_bytes = f.read()

    st.download_button(
        label="Download file",
        data=pdf_bytes,
        file_name=os.path.basename(pdf_path),
        mime="application/docx"
    )

# Replace 'path/to/your/pdf.pdf' with the actual path to your PDF file
pdf_path = "Summary.docx"


# Create a Streamlit button to trigger the download
if st.button("REsearch"):
    data = parse(topic)

    err = []


    for i in data.keys():
        try:
            a = Sichub2(i)
            down.append(a)
        except:
            err.append(i)

    for i in err:
        try:
            b = google(i,data[i])
            down.append(b)
        except:
            print(i)


    for i in down:
        reader = PdfReader(i)
        print('my head')
        try:
          
          rag_tool = ragtool(i)
          print('my nech')
          words = agents()
          print('my hand')
          you.append(words)
        except Exception as exc:
            print(exc)




    doc()
    download_pdf(pdf_path)


    


#st.download_button('DOwnload button',data = 'h.csv', file_name = 'Summary.csv', mime ="text/docx")
# Use Local CSS File
def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)


local_css("style/style.css")





